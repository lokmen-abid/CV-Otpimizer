# python
"""parsers/cv_parser.py

Parser de CV pour ATS-Optimizer.
Fournit extraction (PDF/DOCX/texte), nettoyage, détection de langue (fr/en)
et détection des sections (regex + spaCy optionnel).
"""
from __future__ import annotations

import logging
import os
import re
import unicodedata
from io import BytesIO
from typing import Union, Dict, Optional, List

logger = logging.getLogger(__name__)

# ----- Helpers pour lecture -----

def _read_bytes_if_needed(file: Union[str, BytesIO, bytes]) -> tuple[Optional[str], Optional[BytesIO]]:
    """Normalise l'entrée en (path, BytesIO|None)."""
    if isinstance(file, (bytes, bytearray)):
        return None, BytesIO(bytes(file))
    if hasattr(file, "read"):
        try:
            data = file.read()
            if isinstance(data, str):
                data = data.encode("utf-8")
            return None, BytesIO(data)
        except Exception as exc:
            logger.exception("Impossible de lire le flux: %s", exc)
            return None, None
    if isinstance(file, str):
        return file, None
    return None, None

# ----- Imports paresseux -----

def _import_fitz():
    try:
        import fitz  # type: ignore
        return fitz
    except Exception as exc:
        raise ImportError("PyMuPDF (fitz) requis pour extraire le texte des PDF.") from exc

def _import_docx():
    try:
        import docx  # type: ignore
        return docx
    except Exception as exc:
        raise ImportError("python-docx requis pour extraire le texte des DOCX.") from exc

# ----- Extraction -----

def extract_text_from_pdf(file):
    fitz = _import_fitz()
    path, stream = _read_bytes_if_needed(file)

    if path:
        doc = fitz.open(path)
    else:
        doc = fitz.open(stream=stream.getvalue(), filetype="pdf")

    text_pages = []
    for page in doc:
        page_blocks = page.get_text("blocks")
        # Trier les blocs : haut en bas (y0 asc), puis gauche à droite (x0 asc)
        page_blocks.sort(key=lambda b: (b[1], b[0]))  # y0, x0
        texts = [b[4].strip() for b in page_blocks if len(b[4].strip()) > 1]
        text_pages.append("\n".join(texts))
    doc.close()
    return "\n\n".join(text_pages)

def extract_text_from_docx(file: Union[str, BytesIO, bytes]) -> str:
    """Extrait texte d'un .docx. Accepte chemin ou bytes/file-like."""
    try:
        docx = _import_docx()
    except ImportError as exc:
        logger.error(str(exc))
        return ""

    path, stream = _read_bytes_if_needed(file)
    try:
        if path:
            if not os.path.exists(path):
                raise FileNotFoundError(f"Fichier introuvable: {path}")
            document = docx.Document(path)
        elif stream:
            document = docx.Document(stream)
        else:
            logger.error("Aucun chemin ni flux valide fourni à extract_text_from_docx")
            return ""

        paragraphs = [p.text for p in document.paragraphs if p.text]
        # inclure tableaux
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.exception("Erreur extraction DOCX: %s", exc)
        return ""

def extract_text(file: Union[str, BytesIO, bytes], filename: Optional[str] = None) -> str:
    """Détecte le type (via filename/extension) et extrait le texte."""
    # préférer filename/str extension
    ext = None
    if filename and isinstance(filename, str) and "." in filename:
        ext = filename.rsplit(".", 1)[1].lower()
    elif isinstance(file, str) and "." in file:
        ext = file.rsplit(".", 1)[1].lower()

    if ext == "pdf":
        return extract_text_from_pdf(file)
    if ext in ("docx", "doc"):
        return extract_text_from_docx(file)

    # fallback: si c'est un chemin texte, essayer lecture directe
    path, stream = _read_bytes_if_needed(file)
    if path and os.path.exists(path):
        try:
            with open(path, "rb") as f:
                raw = f.read()
            # tenter pdf/docx selon magic simple (PDF header %PDF)
            if raw.startswith(b"%PDF"):
                return extract_text_from_pdf(file)
            # tenter docx (zip) heuristique
            if b"word/" in raw[:1024]:
                return extract_text_from_docx(file)
            # sinon tenter lecture texte
            try:
                return raw.decode("utf-8", errors="ignore")
            except Exception:
                return ""
        except Exception as exc:
            logger.exception("Erreur lecture fichier: %s", exc)
            return ""
    else:
        # tenter via byte/file-like extracteurs
        return extract_text_from_docx(file) or extract_text_from_pdf(file) or ""

# ----- Nettoyage -----

def clean_text(text: Optional[str], collapse_newlines: bool = True) -> str:
    """Nettoie et normalise le texte: unicode, espaces, caractères non imprimables."""
    if not text:
        return ""
    try:
        txt = unicodedata.normalize("NFKC", text)
        txt = txt.replace("\r\n", "\n").replace("\r", "\n")
        txt = txt.replace("•", "-").replace("–", "-").replace("—", "-")  # Normalise bullets/dash
        txt = "".join(ch for ch in txt if ch.isprintable() or ch == "\n")
        txt = re.sub(r"[ \t]{2,}", " ", txt)
        if collapse_newlines:
            txt = re.sub(r"\n{3,}", "\n\n", txt)
        txt = re.sub(r"\n\s+\n+", "\n\n", txt)
        return txt.strip()
    except Exception as exc:
        logger.exception("Erreur lors du nettoyage du texte: %s", exc)
        return text or ""

# ----- Sections / Headers -----

def normalize_header(text: str) -> str:
    text = text.lower()
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r"[^a-z0-9\s]", "", text)  # Enlève & , : etc. pour matching simple
    text = re.sub(r"\s+", " ", text).strip()
    return text

_SECTION_HEADERS: dict[str, List[str]] = {
    "profile": ["profile", "profile summary", "summary", "profil", "résumé", "resume", "objective"],
    "experience": ["experience", "work experience", "expérience", "expériences", "professional experience", "employment history"],
    "education": ["education", "formation", "education & certifications", "academic qualifications", "studies", "éducation", "degrees"],
    "skills": ["skills", "compétences", "technical skills", "skills & tools", "compétence", "abilities", "technical expertise"],
    "languages": ["languages", "langues", "language skills"],
    "certifications": ["certifications", "certificat", "certificates", "certifications & trainings"],
    "projects": ["projects", "academic projects", "projets", "personal projects", "portfolio"],
    "contact": ["contact", "contactez", "informations personnelles", "personal info", "contact information"],
    "strengths": ["strengths & qualities", "strengths", "qualities", "personal qualities", "atouts", "leadership & community", "leadership and community"],
}

_NORM_HEADER_TO_SECTION: dict[str, str] = {}
for _key, _variants in _SECTION_HEADERS.items():
    for _v in _variants:
        _norm = normalize_header(_v)
        _NORM_HEADER_TO_SECTION.setdefault(_norm, _key)

def _looks_like_contact_line(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    sl = s.lower()
    if "@" in s and re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", s, flags=re.IGNORECASE):
        return True
    if any(k in sl for k in ["linkedin", "github", "portfolio", "website", "www.", "http://", "https://"]):
        return True
    # Téléphone (heuristique simple): au moins 8 chiffres au total
    digits = re.sub(r"\D", "", s)
    if len(digits) >= 8:
        return True
    return False

def _looks_like_name_line(line: str) -> bool:
    s = line.strip()
    if not s or any(ch.isdigit() for ch in s):
        return False
    if "@" in s or "http" in s.lower():
        return False
    words = s.split()
    if not (1 <= len(words) <= 4):
        return False
    letters = re.sub(r"[^A-Za-zÀ-ÖØ-öø-ÿ]", "", s)
    if len(letters) < 3:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
    return upper_ratio >= 0.8

def is_potential_header(line: str) -> bool:
    """Vérifie si une ligne est un potentiel header: court, majuscules, ou finit par : ."""
    stripped = line.strip()
    if not stripped:
        return False
    words = stripped.split()
    if len(words) > 6 or len(words) < 1:
        return False
    # Si c'est un header connu (même en Title Case), accepter
    if normalize_header(stripped) in _NORM_HEADER_TO_SECTION:
        return True
    # Ignorer si ressemble à un sous-titre (ex. "Data Analysis & Business Intelligence")
    if len(words) > 3 and not stripped.upper() == stripped:
        return False
    if stripped.upper() == stripped or stripped.endswith((":","–","-")):
        return True
    return False

def split_sections_by_headers(text: str) -> Dict[str, str]:
    """Split ligne par ligne pour détecter headers et collecter contenu."""
    sections: Dict[str, List[str]] = {}
    current_section: Optional[str] = None
    seen_any_header = False
    lines = text.splitlines()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue  # Ignorer lignes vides

        norm = normalize_header(stripped)
        matched_section = None

        # Match exact sur les headers connus (même sans ':' / majuscules)
        if norm in _NORM_HEADER_TO_SECTION:
            matched_section = _NORM_HEADER_TO_SECTION[norm]
        elif is_potential_header(line):
            # Match plus flexible si la ligne ressemble à un header
            for key, variants in _SECTION_HEADERS.items():
                for v in variants:
                    v_norm = normalize_header(v)
                    if norm == v_norm:
                        matched_section = key
                        break
                    if norm.startswith(v_norm + " "):
                        extra_words = norm[len(v_norm):].strip().split()
                        if 1 <= len(extra_words) <= 2:
                            matched_section = key
                            break
                if matched_section:
                    break

        if matched_section:
            current_section = matched_section
            seen_any_header = True
            continue  # Ne pas ajouter la ligne header au contenu

        if current_section is None:
            current_section = "contact" if (_looks_like_contact_line(stripped) or _looks_like_name_line(stripped)) else "profile"
        elif current_section == "contact" and not seen_any_header:
            # Ne pas engloutir tout le CV dans "contact" si on dépasse le bloc d'en-tête
            if not (_looks_like_contact_line(stripped) or _looks_like_name_line(stripped)):
                if sections.get("contact"):
                    current_section = "profile"

        sections.setdefault(current_section, [])
        sections[current_section].append(line)

    # Convertir listes en str
    result = {k: "\n".join(v).strip() for k, v in sections.items()}

    # Fallback si vide
    if not result:
        result["unknown"] = text

    return result

# ----- Détection langue -----

def detect_language(text: str) -> str:
    """Heuristique simple pour déterminer 'fr' ou 'en'."""
    if not text:
        return "en"
    sample = text[:5000].lower()
    fr_score = sum(sample.count(w) for w in [" le ", " la ", " et ", " de ", " à ", "les ", "des ", "profil", "compétences", "expérience", "formation"])
    en_score = sum(sample.count(w) for w in [" the ", " and ", " of ", " in ", " to ", " skills ", " experience ", " education ", " summary "])
    return "fr" if fr_score > en_score else "en"

# ----- Détection des sections avec spaCy optionnel -----

def detect_sections(text: str, lang: Optional[str] = None) -> Dict[str, str]:
    """Detecte des sections en combinant regex et spaCy (si disponible)."""
    if not text:
        return {}

    if lang is None:
        lang = detect_language(text)

    sections = split_sections_by_headers(text)

    # Enrichir avec spaCy si possible
    try:
        import spacy  # type: ignore

        model = "fr_core_news_sm" if lang == "fr" else "en_core_web_sm"
        try:
            nlp = spacy.load(model)
        except Exception:
            nlp = spacy.blank("fr" if lang == "fr" else "en")
            if "sentencizer" not in nlp.pipe_names:
                nlp.add_pipe("sentencizer")

        doc = nlp(text)
        for sent in doc.sents:
            s = sent.text.strip()
            if not s or not is_potential_header(s):
                continue
            norm = normalize_header(s)
            matched_key = None
            for k, variants in _SECTION_HEADERS.items():
                for v in variants:
                    if norm == normalize_header(v):
                        matched_key = k
                        break
                if matched_key:
                    break
            if matched_key and not sections.get(matched_key):
                start = text.find(s)
                if start >= 0:
                    after = text[start + len(s):]
                    m = re.search(r"\n\n|\n{2,}", after)  # Trouver fin de section
                    if m:
                        chunk = after[:m.start()].strip()
                    else:
                        chunk = after.strip()
                    if chunk:
                        sections[matched_key] = chunk
    except Exception as exc:
        logger.debug("spaCy indisponible ou erreur: %s", exc)

    # Garantir présence des clés connues
    for k in _SECTION_HEADERS.keys():
        sections.setdefault(k, "")

    return sections

def post_process_sections(sections: Dict[str, str]) -> Dict[str, str]:
    """Post-traitement pour corriger mélanges courants."""
    # Déplacer education de skills si mots-clés
    if "skills" in sections and any(word in sections["skills"].lower() for word in ["university", "bachelor", "diploma", "certification"]):
        sections["education"] = (sections["education"] + "\n\n" + sections["skills"]).strip()
        sections["skills"] = ""

    # Séparer languages d'autres sections si lignes matching (kw langue ET dash)
    lang_keywords = ["french", "english", "arabic", "advanced", "native", "langues", "languages"]
    for src_key in ["education", "strengths", "profile"]:
        if src_key in sections:
            lines = sections[src_key].splitlines()
            src_content, lang_content = [], []
            for l in lines:
                if any(kw in l.lower() for kw in lang_keywords) and ("-" in l or "—" in l):
                    lang_content.append(l)
                else:
                    src_content.append(l)
            sections[src_key] = "\n".join(src_content).strip()
            if lang_content:
                sections["languages"] = (sections.get("languages", "") + "\n" + "\n".join(lang_content)).strip()

    # Merger projects dans experience si petit
    if "projects" in sections and len(sections["projects"].split()) < 50 and "experience" in sections:
        sections["experience"] += "\n\n" + sections["projects"]
        del sections["projects"]

    return sections

# ----- Orchestrateur -----

def parse_cv(file: Union[str, BytesIO, bytes], filename: Optional[str] = None, do_clean: bool = True) -> Dict[str, Union[str, Dict[str, str]]]:
    """Orchestre extraction, nettoyage et détection de sections.

    Retourne dict avec: filename, raw_text, cleaned_text, language, sections.
    En cas d'erreur d'extraction, les champs existent mais peuvent être vides.
    """
    result = {
        "filename": filename or (file if isinstance(file, str) else ""),
        "raw_text": "",
        "cleaned_text": "",
        "language": "",
        "sections": {},
    }

    try:
        raw = extract_text(file, filename=filename)
        result["raw_text"] = raw or ""
        cleaned = clean_text(raw) if do_clean else raw or ""
        result["cleaned_text"] = cleaned
        lang = detect_language(cleaned)
        result["language"] = lang
        sections = detect_sections(cleaned, lang=lang)
        sections = post_process_sections(sections)
        result["sections"] = sections
    except Exception as exc:
        logger.exception("Erreur lors du parsing du CV: %s", exc)
    return result

__all__ = [
    "extract_text_from_pdf",
    "extract_text_from_docx",
    "extract_text",
    "clean_text",
    "detect_language",
    "detect_sections",
    "parse_cv",
]
